# Barbara Makowski
# p2.py

import docx
import os
from docx.shared import Pt



# removes headers and footers
def remove_header_and_footer(doc):
    # removes headers and footers
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = True

        footer = section.footer
        footer.is_linked_to_previous = True

# deletes first empty paragraoh of document
def delete_first_extra_paragraph(doc):
    first_paragraph = doc.paragraphs[0]

    if not first_paragraph.text:
        first_paragraph._element.getparent().remove(first_paragraph._element)  

# to delete the extra returns throughout the document
def delete_extrapara(doc):
    for i in range(len(doc.paragraphs)-1, -1, -1):
        para = doc.paragraphs[i]
        if not para.text.strip():
            doc._element.body.remove(para._element)
        else:
            para_format = para.paragraph_format
            para_format.space_before = docx.shared.Pt(0)
            para_format.space_after = docx.shared.Pt(0)

# removes section breaks
def remove_section_breaks(doc):
    for i in range(len(doc.paragraphs)):
        p = doc.paragraphs[i]
    
    # Check if the paragraph style contains the string "section"
    if "section" in p.style.name.lower():
        
        # Remove the paragraph
        doc._element.body[i].getparent().remove(doc._element.body[i])

"""
Returns a new cleaned file

doc: (docx.document) document logic performed on
"""
def readtxt(doc):
    outputDoc = docx.Document()
    fullText = []

    style = doc.styles['Normal']
    font = style.font
    # when font is in Arial, 'Transcript' and 'L' in FINAL deletes, but not Bloomberg or FINA
    font.name = 'Arial'
    font.size = Pt(10)
    
    # strips to plain text, changes to Calibri font
    for para in doc.paragraphs:
        
        if len(para.text) > 1: 
            fullText.append((para.text) + "\n")
            outputDoc.add_paragraph(para.text + "\n") 

    # until this point, returns all text with headers and footers, only BLOOMBERG on side, and FINA (no L) 

    remove_header_and_footer(doc)

    # until this point, does all from above and removes headers and footers

    delete_first_extra_paragraph(doc)  

    #until this point, does all from above and removes first empty paragraph

    delete_extrapara(doc)

    # until this point, does all from above and removes any extra paragraphs

    remove_section_breaks(doc)

    # until this point, does all from above and removes any loose section breaks; also removed biotags and some of the 
    # bloombergs and finals

    return outputDoc

"""
Loops through folder of files, cleans, and saves in new folder.

Precondition: Folder of documents (must be docx files).
Postcondition: Saves cleaned files to new folder.
"""
def readwrite():
    inputFolder = "MultiFileFolder"
    outputFolder = "CleanedFileFolder"

    for filename in os.listdir(inputFolder):
        
        doc = docx.Document(inputFolder + "/" + filename)
        readtxt(doc)

        outputFilename = "Clean_" + filename
        doc.save(outputFolder + "/" + outputFilename)
        

readwrite()